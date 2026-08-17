"""Build the teacher-template-compatible coursework DOCX from executed results."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT / "coursework/report"
OUTPUT = REPORT_DIR / "实验报告-大作业-组号(1)-HybridSN基线阶段.docx"
TEMPLATE_ASSETS = REPORT_DIR / "template_audit/pdf_images"
FIGURES = PROJECT_ROOT / "coursework/outputs/report_figures"
STAGE1 = PROJECT_ROOT / "coursework/outputs/stage1/pavia_university"
STAGE2 = PROJECT_ROOT / "coursework/outputs/stage2"
STAGE3 = PROJECT_ROOT / "coursework/outputs/stage3/pavia_traditional"
PREPROCESSING = PROJECT_ROOT / "coursework/outputs/comparisons/preprocessing_svm"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def mark_table_header(row) -> None:
    """Mark a semantic table's first row as a repeating header row."""
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_picture_alt(inline_shape, text: str) -> None:
    """Attach meaningful alternative text to an inline Word picture."""
    inline_shape._inline.docPr.set("descr", text)
    inline_shape._inline.docPr.set("title", text)


def set_run_font(run, *, name="宋体", size=10.5, bold=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    set_run_font(run, size=9)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.3)
    section.footer_distance = Cm(1.3)
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(0)
    for style_name, size in (("Title", 26), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = document.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
    document.styles["Heading 1"].paragraph_format.page_break_before = True
    caption = document.styles["Caption"]
    caption.font.name = "宋体"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.keep_with_next = True


def add_text(document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullets(document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.37)
        paragraph.paragraph_format.line_spacing = 1.35
        set_run_font(paragraph.add_run(item))


def add_numbered(document, items: list[str]) -> None:
    for number, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.37)
        paragraph.paragraph_format.line_spacing = 1.35
        set_run_font(paragraph.add_run(f"{number}. {item}"))


def add_table(document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    mark_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(header), size=9, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            set_run_font(paragraph.add_run(str(value)), size=8.5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(document, path: Path, caption: str, *, width=6.25) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.keep_with_next = True
    picture = paragraph.add_run().add_picture(str(path), width=Inches(width))
    set_picture_alt(picture, caption)
    cap = document.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    set_run_font(cap.add_run(caption), size=9)


def add_code_block(document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(lines):
        run = paragraph.add_run(line + ("\n" if index < len(lines) - 1 else ""))
        set_run_font(run, name="Consolas", size=8.5)


def add_cover(document: Document) -> None:
    for _ in range(2):
        document.add_paragraph()
    logo_table = document.add_table(rows=1, cols=2)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_table.autofit = False
    for cell in logo_table.rows[0].cells:
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    seal = logo_table.cell(0, 0).paragraphs[0].add_run().add_picture(
        str(TEMPLATE_ASSETS / "image-000.png"), width=Inches(1.05)
    )
    set_picture_alt(seal, "北京理工大学校徽")
    logotype = logo_table.cell(0, 1).paragraphs[0].add_run().add_picture(
        str(TEMPLATE_ASSETS / "image-001.png"), width=Inches(2.85)
    )
    set_picture_alt(logotype, "北京理工大学中英文校名标识")
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    set_run_font(title.add_run("本科实验报告"), name="黑体", size=28, bold=True)
    for _ in range(2):
        document.add_paragraph()
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.first_line_indent = Cm(0)
    set_run_font(name.add_run("实验名称："), name="黑体", size=15, bold=True)
    experiment_run = name.add_run(" 大作业5：高光谱智能解译 ")
    set_run_font(experiment_run, size=15)
    experiment_run.underline = True

    table = document.add_table(rows=8, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    column_widths = (Inches(1.25), Inches(2.30), Inches(1.25), Inches(1.45))
    for row in table.rows:
        for cell, width in zip(row.cells, column_widths, strict=True):
            cell.width = width
    labels = [
        ("课程名称：", "人工智能基础理论与技术\n人工智能导论", "实验时间：", "课后"),
        ("任课教师：", "刘欢", "实验地点：", ""),
        ("助　教　：", "阮子航", "实验类型：", "□ 原理验证\n■ 综合设计\n□ 自主创新"),
        ("姓名(组长)：", "", "", ""),
        ("学号(组长)：", "", "学院(组长)：", ""),
        ("组　员　：", "", "", ""),
        ("组　号　：", "", "成　绩：", ""),
        ("提交阶段：", "基线、消融与传统方法对比", "随机种子：", "1442"),
    ]
    # Merge member row and keep the template's compact matrix.
    table.cell(5, 1).merge(table.cell(5, 3))
    for row_index, values in enumerate(labels):
        cells = table.rows[row_index].cells
        usable = values if row_index != 5 else (values[0], values[1])
        for col_index, value in enumerate(usable):
            cell = cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=90, start=90, bottom=90, end=90)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(paragraph.add_run(value), size=10.5)
        if row_index != 5:
            for col_index in range(len(usable), 4):
                paragraph = cells[col_index].paragraphs[0]
                paragraph.paragraph_format.first_line_indent = Cm(0)
    document.add_page_break()


def build() -> Path:
    document = Document()
    configure_document(document)
    add_cover(document)

    document.add_heading("一、实验目的", level=1)
    document.add_heading("（一）基本任务", level=2)
    add_numbered(
        document,
        [
            "掌握 Pavia University、Indian Pines 与 Salinas 高光谱立方体和标签图的读取、统计、可视化与固定样本构建方法。",
            "学习原始全波段、PCA、LDA、均匀波段选择和 Fisher 波段选择，并保证所有统计变换只在训练集上拟合。",
            "使用 PyTorch 实现 HybridSN，完成训练、验证、测试、性能计时、混淆矩阵、逐类准确率与分类图输出。",
        ],
    )
    document.add_heading("（二）进阶任务", level=2)
    add_numbered(
        document,
        [
            "建立 YAML 调参接口，使卷积通道、卷积核、全连接层、Dropout、BatchNorm、优化器和 epoch 可独立修改。",
            "比较 Softmax+交叉熵与 Sigmoid+one-vs-rest BCE 分类目标。",
            "复现 PCA/LDA/选带、LBP/Gabor、SVM/XGBoost 传统路径并比较精度和效率。",
        ],
    )
    document.add_heading("（三）提高任务", level=2)
    add_text(
        document,
        "形成三阶段可复用工程，支持三数据集；审计随机像元划分的空间重叠，为空间块划分、多 seed 重复、跨场景泛化、注意力与轻量化改进提供可信基线。",
    )

    document.add_heading("二、实验环境", level=1)
    document.add_heading("2.1 数据集", level=2)
    add_table(
        document,
        ["数据集", "立方体尺寸", "波段", "类别", "有标签像元", "状态"],
        [
            ["Pavia University", "610×340×103", "103", "9", "42,776", "完整训练与对比"],
            ["Indian Pines", "145×145×200", "200", "16", "10,249", "已适配/冻结划分"],
            ["Salinas", "512×217×204", "204", "16", "54,129", "已适配/冻结划分"],
        ],
    )
    add_text(
        document,
        "Indian Pines 的 corrected 版本已经去除水吸收波段，程序实际读取 200 个有效波段。三套数据均通过统一注册表声明文件名、MAT 键名和类别名。",
    )
    add_figure(document, STAGE1 / "figures/01_dataset_overview.png", "图1　Pavia University 伪彩色图与标签图")
    document.add_heading("2.2 软件与运行环境", level=2)
    add_bullets(
        document,
        [
            "Python 3.12；PyTorch、NumPy、SciPy、scikit-learn、Matplotlib、pandas、PyYAML。",
            "XGBoost 3.4.1；Notebook 使用项目 .venv 中的 hsi-coursework 内核。",
            "深度模型设备由 runtime.device: auto 自动选择，设备、参数量、耗时、吞吐量和模型大小均写入 performance.json。",
        ],
    )
    document.add_heading("2.3 固定实验协议", level=2)
    add_text(
        document,
        "所有实验统一使用 seed=1442。先按类别分层抽取 70% 测试集，再把 30% 建模池按 80%/20% 分为训练和验证，最终比例为 24%/6%/70%。Pavia 实际样本数为训练 10,265、验证 2,567、测试 29,944。三组互斥并覆盖全部有标签像元，验证和测试从不参与标准化、PCA、LDA 或 Fisher 选带的拟合。",
    )
    add_figure(document, STAGE1 / "figures/02_class_and_split_distribution.png", "图2　按类别分层的固定划分及类别不均衡")

    document.add_heading("三、实验原理", level=1)
    document.add_heading("3.1 任务定义", level=2)
    add_text(
        document,
        "高光谱图像是 H×W×B 数据立方体。二维坐标对应像元，长度为 B 的向量描述该位置的光谱响应；标签图中 0 是未标注背景，1…C 是地物类别。",
    )
    add_text(
        document,
        "本实验不是标准端到端语义分割。传统模型读取中心像元光谱和人工邻域特征；HybridSN 读取以有标签像元为中心的 25×25×B patch，并只预测中心像元类别。逐像元滑动后可形成类似语义分割的分类图，因此更准确的名称是高光谱像元分类或基于 patch 的稠密分类。",
    )
    document.add_heading("3.2 研究路线", level=2)
    add_figure(document, FIGURES / "research_architecture.png", "图3　降维、空间特征和分类器的分组研究架构")
    document.add_heading("3.3 光谱与空间预处理", level=2)
    add_text(document, "逐波段标准化为 z=(x-μ)/σ，其中 μ、σ 仅由训练中心像元估计。PCA 最大化投影方差；LDA 利用训练标签最大化类间散度与类内散度之比，维度上限为 C-1。均匀选带不使用标签；Fisher 选带逐波段计算类间/类内散度并选择最高分波段。")
    add_text(document, "LBP 用 8 邻域比较生成局部二值模式，在 9×9 窗口统计 16 bin 直方图；Gabor 使用 4 个方向、2 个频率的复数滤波器，对幅值响应在局部窗口统计均值和标准差。前三个 PCA 分量用于空间特征，最终与 15 维中心光谱拼接。")
    document.add_heading("3.4 HybridSN", level=2)
    add_text(document, "基线输入为 N×1×15×25×25。三层有效 3D 卷积的输出通道为 8、16、32，光谱核深度为 7、5、3，空间核为 3×3；随后合并特征通道与剩余光谱深度，使用 64 通道 2D 卷积，接 256、128 单元全连接层和 Dropout=0.4。输出 9 维 logits，总参数量 4,844,793。")
    add_text(document, "Softmax 主实验采用交叉熵；Sigmoid 消融采用 one-hot 标签的 BCEWithLogitsLoss。两者预测均取最大 logit。单标签互斥分类中 Softmax 的概率语义更自然，Sigmoid 仅作为消融观察。")
    document.add_heading("3.5 评价指标", level=2)
    add_text(document, "总体准确率 OA 为正确样本数除以测试样本数；逐类准确率是混淆矩阵对角元素除以该类支持数；AA 是逐类准确率宏平均；Kappa 校正偶然一致性。另记录训练时间、测试推理时间、吞吐量、参数量和模型文件大小。")

    document.add_heading("四、实验内容和结果分析", level=1)
    document.add_heading("4.1 工程结构与阶段交接", level=2)
    add_text(document, "coursework/ 中包含三个编号 Notebook、分阶段 YAML、输出和报告；src/coursework/ 存放核心。阶段一清单保存数据集、划分、配置指纹、预处理状态和模型就绪 NPZ；阶段二、三只读取该清单，不重新抽样。")
    add_code_block(
        document,
        [
            "coursework/",
            "  configs/{stage1_data, stage2_hybridsn, stage3_traditional}/",
            "  notebooks/01_...  02_...  03_...",
            "  outputs/{stage1, stage2, stage3, comparisons, report_figures}/",
            "src/coursework/{stage1.py, stage2.py, stage3.py, reporting.py}",
        ],
    )
    document.add_heading("4.2 降维和选带控制变量", level=2)
    add_table(
        document,
        ["路线", "维度", "OA/%", "AA/%", "Kappa", "训练/s", "全标注推理/s"],
        [
            ["原始全波段", "103", "94.0556", "91.8837", "0.920732", "12.88", "61.75"],
            ["PCA15", "15", "93.7016", "91.3993", "0.916289", "24.10", "61.23"],
            ["LDA8", "8", "91.3438", "88.8494", "0.884856", "11.93", "10.54"],
            ["均匀15", "15", "90.3620", "87.4884", "0.870633", "1.32", "8.72"],
            ["Fisher15", "15", "70.5884", "56.2779", "0.583684", "3.64", "16.97"],
        ],
    )
    add_figure(
        document,
        PREPROCESSING / "preprocessing_svm_comparison.png",
        "图4　相同 RBF-SVM 下的预处理准确率与效率",
        width=5.4,
    )
    add_text(document, "原始全波段中心光谱 OA 略高于 PCA15，但维度和存储成本更大；PCA 的核心价值是压缩并为 HybridSN 固定输入维度。Fisher15 是明显失败路线，说明逐波段判别性最高不等于波段子集联合最优。后续应尝试 mRMR、互信息、SPA、稀疏约束或可学习光谱注意力。")

    document.add_heading("4.3 传统特征与分类器", level=2)
    add_table(
        document,
        ["方法", "维度", "OA/%", "AA/%", "Kappa", "训练/s", "测试/s"],
        [
            ["PCA15+SVM", "15", "93.7016", "91.3993", "0.916289", "19.02", "34.44"],
            ["PCA15+XGBoost", "15", "92.9201", "90.0874", "0.905551", "15.27", "0.389"],
            ["PCA15+LBP+SVM", "63", "99.5124", "99.3337", "0.993531", "38.22", "63.27"],
            ["PCA15+Gabor+SVM", "63", "99.5692", "99.5332", "0.994291", "22.98", "34.39"],
            ["PCA15+LBP+Gabor+SVM", "111", "99.9365", "99.9212", "0.999159", "33.32", "55.74"],
            ["融合+XGBoost", "111", "99.2052", "98.7912", "0.989464", "28.78", "0.417"],
        ],
    )
    add_figure(document, STAGE3 / "traditional_method_comparison.png", "图5　LBP/Gabor 与 SVM/XGBoost 的传统方法对比")
    add_text(document, "LBP/Gabor 使 OA 从 93.70% 提升到 99.5%以上，说明当前随机像元协议下邻域纹理是主要增益来源。融合 SVM 精度最高但测试吞吐仅约 537 样本/秒；融合 XGBoost OA 为 99.21%，吞吐约 71,728 样本/秒，是更强的效率方案。")
    add_figure(document, STAGE3 / "PCA15_LBP_Gabor_SVM/classification_maps.png", "图6　传统融合 SVM 的真值、测试预测与全标注分类图")

    document.add_page_break()
    document.add_heading("4.4 HybridSN 基线与分类目标消融", level=2)
    add_table(
        document,
        ["目标", "最佳epoch", "验证OA/%", "测试OA/%", "AA/%", "Kappa", "错误", "训练/s", "测试/s"],
        [
            ["Softmax+CE", "5", "99.9610", "99.9065", "99.8353", "0.998761", "28", "110.24", "6.94"],
            ["Sigmoid+BCE", "12", "99.9610", "99.9699", "99.9205", "0.999602", "9", "157.21", "7.06"],
        ],
    )
    add_figure(document, FIGURES / "hybridsn_objective_comparison.png", "图7　Softmax 与 Sigmoid 分类目标的精度和训练成本")
    add_figure(document, STAGE2 / "pavia_softmax_baseline/learning_curves.png", "图8　Softmax 基线训练/验证损失与准确率")
    add_figure(document, STAGE2 / "pavia_sigmoid_ablation/learning_curves.png", "图9　Sigmoid 消融训练/验证损失与准确率")
    add_text(document, "本次单次划分中 Sigmoid 消融比 Softmax 少 19 个错误，但训练时间增加约 42.6%。两者验证 OA 相同且测试只进行一次，不能把小差异解释为稳定提升；应在多个 seed 和空间块划分上复验，并增加 ECE/Brier 等概率校准指标。")
    add_figure(document, STAGE2 / "pavia_softmax_baseline/confusion_matrix.png", "图10　Softmax 基线测试混淆矩阵")
    add_figure(document, STAGE2 / "pavia_sigmoid_ablation/per_class_accuracy.png", "图11　Sigmoid 消融逐类测试准确率")
    add_figure(document, STAGE2 / "pavia_sigmoid_ablation/classification_maps.png", "图12　Sigmoid 消融的真值、测试预测与全标注分类图")

    document.add_heading("4.5 统一精度—效率比较", level=2)
    add_figure(document, FIGURES / "all_method_accuracy_comparison.png", "图13　传统与深度方法的 OA、AA、Kappa 统一比较")
    add_figure(document, FIGURES / "accuracy_efficiency_tradeoff.png", "图14　测试精度、推理时间与模型大小权衡")
    add_text(document, "融合传统 SVM 的 OA 介于两组 HybridSN 之间；HybridSN 测试推理约 7 秒，比融合 RBF-SVM 快约 8 倍，但训练时间更长。XGBoost 最快但精度有损失，形成可根据部署约束选择的 Pareto 方案。")

    document.add_heading("4.6 有效性威胁与改进方向", level=2)
    add_text(document, "25×25 patch 半径为 12。空间重叠审计显示 29,944 个测试 patch 中，29,939 个（99.9833%）至少包含一个训练中心，29,934 个（99.9666%）包含同类训练中心；每个测试 patch 内训练中心中位数为 73。这不是测试标签进入训练的代码泄漏，但说明随机像元划分具有强空间依赖，极高 OA 不能直接代表对新区域的泛化。")
    add_numbered(
        document,
        [
            "使用空间块、连通区域或跨场景划分，避免邻域高度重叠。",
            "至少 5 个 seed 重复并报告均值±标准差和置信区间。",
            "在 Indian Pines、Salinas 完成完整训练并补充跨数据集迁移。",
            "研究轻量 HybridSN、BatchNorm/残差、光谱—空间注意力、Mixup/CutMix 与类别不均衡损失。",
            "同时报告精度、概率校准、参数量、显存、延迟和吞吐量。",
        ],
    )

    document.add_heading("五、参考文献", level=1)
    references = [
        "[1] S. K. Roy et al. HybridSN: Exploring 3-D–2-D CNN Feature Hierarchy for Hyperspectral Image Classification. IEEE GRSL, 2020.",
        "[2] T. Ojala, M. Pietikäinen, T. Mäenpää. Multiresolution Gray-Scale and Rotation Invariant Texture Classification with Local Binary Patterns. IEEE TPAMI, 2002.",
        "[3] C. Cortes, V. Vapnik. Support-Vector Networks. Machine Learning, 1995.",
        "[4] T. Chen, C. Guestrin. XGBoost: A Scalable Tree Boosting System. KDD, 2016.",
        "[5] F. Pedregosa et al. Scikit-learn: Machine Learning in Python. JMLR, 2011.",
        "[6] A. Paszke et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library. NeurIPS, 2019.",
        "[7] 课程参考仓库 High_spectrum_BIT-main：Highspectrum.ipynb、HybridSN.ipynb 及 figures。",
        "[8] University of the Basque Country. Hyperspectral Remote Sensing Scenes datasets.",
    ]
    for reference in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.25
        set_run_font(paragraph.add_run(reference), size=9.5)

    document.add_heading("六、心得体会", level=1)
    add_text(document, "本实验的关键不只是把 HybridSN 跑到较高准确率，而是建立可核查的实验边界。固定 seed 和划分、训练集拟合预处理、验证集选模型、测试集只报告，使模型优化有统一基准。传统融合方法的强表现说明空间纹理对当前随机像元协议影响巨大；Fisher 选带失败也提醒我们不能从方法名称推断效果。后续工作应把注意力从继续追逐 99.9% 转向更严格的空间泛化、多次重复和效率权衡。")

    document.add_heading("附录 A　关键代码与复现命令", level=1)
    add_code_block(
        document,
        [
            ".\\.venv\\Scripts\\python.exe scripts\\运行课程数据预处理.py",
            ".\\.venv\\Scripts\\python.exe scripts\\运行课程HybridSN.py --config ...pavia_softmax_baseline.yaml --output-dir ...pavia_softmax_baseline",
            ".\\.venv\\Scripts\\python.exe scripts\\运行课程传统分类.py",
            ".\\.venv\\Scripts\\python.exe scripts\\execute_coursework_notebooks.py",
        ],
    )
    add_text(document, "三本 Notebook、YAML、模型、CSV、PNG、固定划分、预处理状态和测试预测均随项目提交。")

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
